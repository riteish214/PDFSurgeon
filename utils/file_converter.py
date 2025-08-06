import os
import tempfile
from docx import Document
# Note: docx2pdf requires LibreOffice, using reportlab fallback
import pandas as pd
from pptx import Presentation
import pdfplumber
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
import logging

class FileConverter:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def convert_file(self, input_path, conversion_type):
        """Convert file based on conversion type"""
        try:
            if conversion_type == 'pdf_to_docx':
                return self.pdf_to_docx(input_path)
            elif conversion_type == 'pdf_to_txt':
                return self.pdf_to_txt(input_path)
            elif conversion_type == 'pdf_to_csv':
                return self.pdf_to_csv(input_path)
            elif conversion_type == 'docx_to_pdf':
                return self.docx_to_pdf(input_path)
            elif conversion_type == 'txt_to_pdf':
                return self.txt_to_pdf(input_path)
            elif conversion_type == 'csv_to_pdf':
                return self.csv_to_pdf(input_path)
            else:
                raise Exception(f"Unsupported conversion type: {conversion_type}")
        
        except Exception as e:
            self.logger.error(f"Error converting file: {str(e)}")
            raise Exception(f"Failed to convert file: {str(e)}")
    
    def pdf_to_docx(self, pdf_path):
        """Convert PDF to DOCX using text extraction"""
        try:
            # Extract text from PDF
            text_content = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text_content += page.extract_text() or ""
                    text_content += "\n\n"
            
            # Create DOCX document
            doc = Document()
            doc.add_heading('Converted from PDF', 0)
            
            # Split text into paragraphs and add to document
            paragraphs = text_content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            output_path = os.path.join(tempfile.gettempdir(), 'converted.docx')
            doc.save(output_path)
            
            return output_path
        
        except Exception as e:
            self.logger.error(f"Error converting PDF to DOCX: {str(e)}")
            raise Exception(f"Failed to convert PDF to DOCX: {str(e)}")
    
    def pdf_to_txt(self, pdf_path):
        """Convert PDF to TXT"""
        try:
            text_content = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text_content += page.extract_text() or ""
                    text_content += "\n\n"
            
            output_path = os.path.join(tempfile.gettempdir(), 'converted.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            return output_path
        
        except Exception as e:
            self.logger.error(f"Error converting PDF to TXT: {str(e)}")
            raise Exception(f"Failed to convert PDF to TXT: {str(e)}")
    
    def pdf_to_csv(self, pdf_path):
        """Extract tables from PDF and convert to CSV"""
        try:
            tables = []
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
            
            if not tables:
                raise Exception("No tables found in PDF")
            
            # Convert first table to CSV
            df = pd.DataFrame(tables[0][1:], columns=tables[0][0])
            output_path = os.path.join(tempfile.gettempdir(), 'converted.csv')
            df.to_csv(output_path, index=False)
            
            return output_path
        
        except Exception as e:
            self.logger.error(f"Error converting PDF to CSV: {str(e)}")
            raise Exception(f"Failed to convert PDF to CSV: {str(e)}")
    
    def docx_to_pdf(self, docx_path):
        """Convert DOCX to PDF using reportlab (fallback method)"""
        try:
            # Read DOCX content
            doc = Document(docx_path)
            
            # Extract text content
            text_content = []
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Create PDF using reportlab
            output_path = os.path.join(tempfile.gettempdir(), 'converted.pdf')
            doc_pdf = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            story = []
            for text in text_content:
                if text.strip():
                    p = Paragraph(text, styles['Normal'])
                    story.append(p)
            
            doc_pdf.build(story)
            return output_path
        
        except Exception as e:
            self.logger.error(f"Error converting DOCX to PDF: {str(e)}")
            raise Exception(f"Failed to convert DOCX to PDF: {str(e)}")
    
    def txt_to_pdf(self, txt_path):
        """Convert TXT to PDF"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            output_path = os.path.join(tempfile.gettempdir(), 'converted.pdf')
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            story = []
            paragraphs = text_content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = Paragraph(paragraph.strip(), styles['Normal'])
                    story.append(p)
            
            doc.build(story)
            return output_path
        
        except Exception as e:
            self.logger.error(f"Error converting TXT to PDF: {str(e)}")
            raise Exception(f"Failed to convert TXT to PDF: {str(e)}")
    
    def csv_to_pdf(self, csv_path):
        """Convert CSV to PDF table"""
        try:
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib import colors
            
            # Read CSV
            df = pd.read_csv(csv_path)
            
            # Prepare data for table
            data = [df.columns.tolist()] + df.values.tolist()
            
            output_path = os.path.join(tempfile.gettempdir(), 'converted.pdf')
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            
            # Create table
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story = [table]
            doc.build(story)
            
            return output_path
        
        except Exception as e:
            self.logger.error(f"Error converting CSV to PDF: {str(e)}")
            raise Exception(f"Failed to convert CSV to PDF: {str(e)}")